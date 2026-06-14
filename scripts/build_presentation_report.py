from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "presentation"
DOCX_PATH = OUT_DIR / "RTL_LLM_Benchmark_Report_June_2026.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8F0F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "1F7A4D"
GOLD = "9A6700"
RED = "A61B1B"
WHITE = "FFFFFF"

MODELS = ["Qwen3.6 27B", "Qwen3.6 35B-A3B", "Qwen3 Coder", "Qwen2.5 Coder", "DeepSeek Coder"]
INTERNAL_IDS = [
    "qwen36-27b",
    "qwen36-35b-a3b",
    "qwen3-coder-30b-a3b-instruct",
    "qwen25-coder-32b",
    "deepseek-coder-v2-lite-instruct",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=10.5, bold=False, color="000000", italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_font(run, size=9, color=MID_GRAY)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, label_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_font(r, size=11, bold=True, color=label_color)
    r = p.add_run(text)
    set_font(r, size=11)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_data_table(doc, headers, rows, widths, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=9, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], LIGHT_GRAY)
            p = cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_font(r, size=8.8)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_picture_with_alt(doc, path, width, alt_text):
    shape = doc.add_picture(str(path), width=width)
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)
    return shape


def build_charts():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    font_dir = Path("C:/Windows/Fonts")
    regular = ImageFont.truetype(str(font_dir / "arial.ttf"), 24)
    small = ImageFont.truetype(str(font_dir / "arial.ttf"), 20)
    bold = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 28)

    def chart(path, title, series, ymax, colors):
        image = Image.new("RGB", (1500, 650), "white")
        draw = ImageDraw.Draw(image)
        left, top, right, bottom = 110, 85, 1440, 515
        draw.text((left, 22), title, font=bold, fill="#17365D")
        for step in range(5):
            value = ymax * step / 4
            y = bottom - (bottom - top) * step / 4
            draw.line((left, y, right, y), fill="#D8DEE8", width=2)
            draw.text((22, y - 12), f"{value:.1f}", font=small, fill="#667085")
        draw.line((left, top, left, bottom), fill="#667085", width=2)
        draw.line((left, bottom, right, bottom), fill="#667085", width=2)
        group_w = (right - left) / len(MODELS)
        bar_w = min(64, group_w / (len(series) + 1))
        labels = ["Qwen3.6\n27B", "Qwen3.6\n35B-A3B", "Qwen3\nCoder", "Qwen2.5\nCoder", "DeepSeek\nCoder"]
        for idx, label in enumerate(labels):
            cx = left + group_w * (idx + 0.5)
            total_w = bar_w * len(series)
            for s_idx, (_, values) in enumerate(series):
                value = values[idx]
                x0 = cx - total_w / 2 + s_idx * bar_w
                x1 = x0 + bar_w * 0.8
                y0 = bottom - (bottom - top) * value / ymax
                draw.rectangle((x0, y0, x1, bottom), fill=colors[s_idx])
                draw.text((x0 - 1, y0 - 27), f"{value:.3f}", font=small, fill="#344054")
            lines = label.split("\n")
            for line_idx, line in enumerate(lines):
                box = draw.textbbox((0, 0), line, font=small)
                draw.text((cx - (box[2] - box[0]) / 2, bottom + 15 + line_idx * 23), line, font=small, fill="#344054")
        legend_x = right - 310
        for idx, (name, _) in enumerate(series):
            y = 32 + idx * 28
            draw.rectangle((legend_x, y, legend_x + 22, y + 18), fill=colors[idx])
            draw.text((legend_x + 31, y - 4), name, font=small, fill="#344054")
        image.save(path)

    ve_path = OUT_DIR / "verilogeval_chart.png"
    chart(
        ve_path,
        "VerilogEval v2 functional correctness",
        [("pass@1", [0.6154, 0.5705, 0.4808, 0.4487, 0.4872]), ("pass@5", [0.7756, 0.7308, 0.5705, 0.5385, 0.5641])],
        0.9,
        ["#2E74B5", "#70AD47"],
    )
    rtl_path = OUT_DIR / "rtlopt_chart.png"
    chart(
        rtl_path,
        "RTL-OPT behavior-preserving equivalence",
        [("equivalence pass", [0.625, 0.475, 0.475, 0.475, 0.725])],
        0.8,
        ["#2E74B5"],
    )
    return ve_path, rtl_path


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("SILICONCRAFT  |  RTL LLM BENCHMARK")
    set_font(r, size=8.5, bold=True, color=MID_GRAY)
    footer = section.footer
    p = footer.paragraphs[0]
    r = p.add_run("Public benchmark briefing  |  June 2026")
    set_font(r, size=8.5, color=MID_GRAY)
    add_page_number(footer.add_paragraph())


def build_doc():
    ve_chart, rtl_chart = build_charts()
    doc = Document()
    configure_document(doc)

    # Memo-style opening block.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("TECHNICAL BENCHMARK REPORT")
    set_font(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Evaluating Large Language Models for RTL Generation")
    set_font(r, size=25, bold=True, color=NAVY, name="Aptos Display")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("A reproducible comparison across VerilogEval v2, RTLLM 2.0, ProtocolLLM, and RTL-OPT")
    set_font(r, size=13, color=MID_GRAY)

    metadata = [
        ("Prepared by", "SiliconCraft RTL Benchmarking Project"),
        ("Report date", "14 June 2026"),
        ("Models evaluated", "Five public LLM baselines"),
        ("Compute", "Lanta GPU cluster, OpenAI-compatible vLLM endpoints"),
        ("Status", "Completed baseline comparison"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_font(r, bold=True, color=NAVY)
        r = p.add_run(value)
        set_font(r)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_callout(
        doc,
        "Bottom line",
        "Qwen3.6 27B is the strongest overall model for functional RTL generation. DeepSeek-Coder-V2-Lite-Instruct leads the RTL-OPT equivalence test, showing the best behavior-preservation rate, but it rarely produces a smaller equivalent circuit.",
    )

    add_heading(doc, "Executive Summary", 1)
    add_body(doc, "We built a reusable benchmark harness that can swap models through an OpenAI-compatible API while keeping prompts, sampling settings, extraction logic, and evaluation tools consistent. Five models were evaluated under matched conditions.")
    add_bullet(doc, "Qwen3.6 27B achieved the best VerilogEval v2 functional pass@1 (61.54%) and pass@5 task recovery (77.56%).")
    add_bullet(doc, "Qwen3.6 27B and Qwen3.6 35B-A3B tied for the best RTLLM 2.0 functional result (60.00%).")
    add_bullet(doc, "DeepSeek Coder achieved the best RTL-OPT equivalence pass rate (72.50%), ahead of Qwen3.6 27B (62.50%).")
    add_bullet(doc, "Qwen2.5 Coder produced the strongest VerilogEval syntax rate, but that did not translate into functional correctness.")
    add_bullet(doc, "ProtocolLLM results are lint-only and must not be interpreted as protocol functional correctness.")

    add_heading(doc, "Recommendation", 2)
    add_body(doc, "Use Qwen3.6 27B as the current default baseline for general RTL generation. Keep DeepSeek Coder as a complementary candidate for optimization workflows where behavior preservation is the first gate. Do not select a model based on syntax or lint scores alone.")

    doc.add_page_break()
    add_heading(doc, "1. What Was Evaluated", 1)
    add_body(doc, "The benchmark suite covers three distinct questions. These measurements are deliberately reported separately because they do not prove the same capability.")
    add_data_table(
        doc,
        ["Benchmark", "Question answered", "Trusted metric", "Scale"],
        [
            ["VerilogEval v2", "Can the model generate functionally correct RTL from a specification?", "Icarus simulation pass@1 / pass@5", "156 tasks"],
            ["RTLLM 2.0", "Can the model solve a second public RTL generation task set?", "Icarus functional pass@1", "50 tasks"],
            ["ProtocolLLM", "Does generated protocol RTL parse and lint?", "Verilator lint pass only", "9 public tasks"],
            ["RTL-OPT", "Can the model rewrite RTL while preserving behavior?", "Yosys equivalence pass", "40 tasks"],
        ],
        [1.15, 2.75, 1.75, 0.85],
    )

    add_heading(doc, "Models", 2)
    for model in MODELS:
        add_bullet(doc, model)

    add_heading(doc, "Controlled Settings", 2)
    add_body(doc, "All model comparisons use the same benchmark version, task set, neutral prompt template, extraction logic, timeout policy, and simulator/evaluation flow. Sampling settings were also matched by benchmark condition.")
    add_data_table(
        doc,
        ["Condition", "Samples/task", "Temperature", "Top-p", "Max tokens"],
        [
            ["VerilogEval pass@1", "1", "0.2", "0.95", "2,048"],
            ["VerilogEval pass@5", "5", "0.6", "0.95", "2,048"],
            ["RTLLM / ProtocolLLM / RTL-OPT", "1", "0.2", "0.95", "4,096"],
        ],
        [2.45, 1.0, 1.0, 0.85, 1.2],
        numeric_cols=(1, 2, 3, 4),
    )

    add_callout(doc, "Interpretation rule", "A lint pass is not a functional pass. A synthesis pass is not proof of behavior preservation. For RTL-OPT, only equivalence-passing outputs provide trustworthy optimization evidence.", fill="FFF4E5", label_color=GOLD)

    doc.add_page_break()
    add_heading(doc, "2. Functional RTL Generation", 1)
    add_body(doc, "Functional simulation is the primary evidence for normal RTL generation. Higher values are better.")
    add_picture_with_alt(
        doc,
        ve_chart,
        Inches(6.35),
        "Grouped bar chart comparing VerilogEval v2 pass@1 and pass@5 functional rates for five RTL language models.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Figure 1. VerilogEval v2 functional pass rates under matched settings.")
    set_font(r, size=8.5, italic=True, color=MID_GRAY)

    add_data_table(
        doc,
        ["Model", "VE pass@1", "VE pass@5", "RTLLM pass@1"],
        [
            ["Qwen3.6 27B", "61.54%", "77.56%", "60.00%"],
            ["Qwen3.6 35B-A3B", "57.05%", "73.08%", "60.00%"],
            ["Qwen3 Coder", "48.08%", "57.05%", "50.00%"],
            ["Qwen2.5 Coder", "44.87%", "53.85%", "54.00%"],
            ["DeepSeek Coder", "48.72%", "56.41%", "52.00%"],
        ],
        [2.45, 1.35, 1.35, 1.35],
        numeric_cols=(1, 2, 3),
    )
    add_heading(doc, "What this means", 2)
    add_bullet(doc, "Qwen3.6 27B is the clearest general-purpose winner across both VerilogEval sampling conditions.")
    add_bullet(doc, "Qwen3.6 35B-A3B is consistently second on VerilogEval and ties the 27B model on RTLLM.")
    add_bullet(doc, "DeepSeek Coder is the best non-Qwen model on VerilogEval pass@1, but Qwen3 Coder slightly overtakes it on pass@5.")
    add_bullet(doc, "The pass@5 gain shows that sampling multiple candidates materially improves task recovery for every model.")

    doc.add_page_break()
    add_heading(doc, "3. Reliability and Failure Patterns", 1)
    add_body(doc, "A model can produce syntactically valid RTL that is still behaviorally wrong. The benchmark therefore tracks extraction, compilation, and simulation failures separately.")
    add_data_table(
        doc,
        ["Model", "VE syntax pass@1", "VE functional pass@1", "Gap"],
        [
            ["Qwen3.6 27B", "83.97%", "61.54%", "22.43 pts"],
            ["Qwen3.6 35B-A3B", "75.64%", "57.05%", "18.59 pts"],
            ["Qwen3 Coder", "86.54%", "48.08%", "38.46 pts"],
            ["Qwen2.5 Coder", "88.46%", "44.87%", "43.59 pts"],
            ["DeepSeek Coder", "85.90%", "48.72%", "37.18 pts"],
        ],
        [2.35, 1.55, 1.55, 1.05],
        numeric_cols=(1, 2, 3),
    )
    add_callout(doc, "Key insight", "The coder-specialized models are very good at producing parseable RTL, but their much lower simulation pass rates reveal semantic errors in timing, state transitions, reset behavior, or protocol logic.")

    add_heading(doc, "VerilogEval pass@1 failure counts", 2)
    add_data_table(
        doc,
        ["Model", "Passed", "Simulation", "Compile", "Extraction"],
        [
            ["Qwen3.6 27B", "96", "35", "17", "8"],
            ["Qwen3.6 35B-A3B", "89", "29", "26", "12"],
            ["Qwen3 Coder", "75", "60", "20", "1"],
            ["Qwen2.5 Coder", "70", "68", "18", "0"],
            ["DeepSeek Coder", "76", "58", "21", "1"],
        ],
        [2.2, 1.0, 1.2, 1.0, 1.1],
        numeric_cols=(1, 2, 3, 4),
    )
    add_body(doc, "The dominant weakness for Qwen3 Coder, Qwen2.5 Coder, and DeepSeek Coder is simulation failure rather than extraction failure. Improving prompt packaging alone is unlikely to close this gap; the models need stronger functional reasoning or an explicitly separate repair loop.")

    doc.add_page_break()
    add_heading(doc, "4. Protocol Lint and RTL Optimization", 1)
    add_heading(doc, "ProtocolLLM public lint", 2)
    add_data_table(
        doc,
        ["Qwen3.6 27B", "Qwen3.6 35B-A3B", "Qwen3 Coder", "Qwen2.5 Coder", "DeepSeek Coder"],
        [["77.78%", "22.22%", "77.78%", "66.67%", "66.67%"]],
        [1.30, 1.30, 1.30, 1.30, 1.30],
        numeric_cols=(0, 1, 2, 3, 4),
    )
    add_callout(doc, "Important limitation", "ProtocolLLM public results are lint-only. They show whether output is structurally acceptable to Verilator, not whether the generated protocol implementation behaves correctly.", fill="FDECEC", label_color=RED)

    add_heading(doc, "RTL-OPT equivalence", 2)
    add_picture_with_alt(
        doc,
        rtl_chart,
        Inches(6.35),
        "Bar chart comparing RTL-OPT equivalence pass rates for five RTL language models; DeepSeek Coder is highest at 0.725.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Figure 2. Fraction of RTL-OPT tasks that synthesized and remained equivalent to the original design.")
    set_font(r, size=8.5, italic=True, color=MID_GRAY)
    add_data_table(
        doc,
        ["Model", "Equiv. pass", "Reduced cells", "Avg. cell ratio"],
        [
            ["Qwen3.6 27B", "25 / 40", "9 / 25", "0.9124"],
            ["Qwen3.6 35B-A3B", "19 / 40", "6 / 19", "0.9003"],
            ["Qwen3 Coder", "19 / 40", "4 / 19", "0.9472"],
            ["Qwen2.5 Coder", "19 / 40", "8 / 19", "0.8709"],
            ["DeepSeek Coder", "29 / 40", "3 / 29", "0.9739"],
        ],
        [2.25, 1.25, 1.35, 1.65],
        numeric_cols=(1, 2, 3),
    )
    add_heading(doc, "Interpretation", 2)
    add_bullet(doc, "DeepSeek Coder preserves behavior most often, passing equivalence on 29 of 40 tasks.")
    add_bullet(doc, "Qwen2.5 Coder has the best average cell ratio among equivalent outputs, but it reaches equivalence on only 19 tasks.")
    add_bullet(doc, "Qwen3.6 27B offers the strongest balance: second-best equivalence coverage and the largest number of equivalence-passing reductions.")
    add_bullet(doc, "A lower cell ratio is only meaningful after equivalence passes; non-equivalent or synthesis-failing outputs are excluded from optimization claims.")

    doc.add_page_break()
    add_heading(doc, "5. Model Selection Guide", 1)
    add_data_table(
        doc,
        ["Use case", "Recommended model", "Why", "Caution"],
        [
            ["General RTL generation", "Qwen3.6 27B", "Best functional pass@1 and pass@5", "Still fails roughly 38% of pass@1 tasks"],
            ["Alternative high-performing baseline", "Qwen3.6 35B-A3B", "Second on VerilogEval; ties RTLLM lead", "Lower syntax and extraction reliability"],
            ["Behavior-preserving RTL rewrite", "DeepSeek Coder", "Best equivalence coverage", "Few equivalent outputs reduce cells"],
            ["Optimization candidate generation", "Qwen3.6 27B / Qwen2.5 Coder", "More valid cell reductions / best ratio", "Use equivalence as a hard gate"],
            ["Syntax-first prototyping", "Qwen2.5 Coder", "Best VerilogEval syntax rate", "Syntax does not imply correct behavior"],
        ],
        [1.45, 1.35, 2.1, 1.6],
    )

    add_heading(doc, "Recommended next experiments", 2)
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run("Analyze simulation failures by semantic category: reset, cycle alignment, FSM transitions, width/sign errors, and handshake logic."))
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run("Add a separately reported repair-loop condition that feeds compiler or simulation errors back to the model."))
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run("Run ProtocolLLM with functional testbenches before making protocol-correctness claims."))
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run("Evaluate optimization quality with technology-mapped area, timing, and power after equivalence passes."))
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run("Repeat selected runs across seeds to quantify variance and confidence intervals."))

    add_heading(doc, "Operational achievement", 2)
    add_body(doc, "The project now has a reusable, model-swappable benchmark workflow. Each run preserves raw responses, extracted RTL, logs, error logs, machine-readable results, summaries, and a per-run report in timestamped folders on Lanta. Committed manifests point to the authoritative output locations without placing large generated artifacts in Git.")

    doc.add_page_break()
    add_heading(doc, "Appendix A. Exact Model IDs", 1)
    add_data_table(doc, ["Report label", "Internal model ID"], list(zip(MODELS, INTERNAL_IDS)), [2.25, 4.25])
    add_heading(doc, "Appendix B. Scope and Limitations", 1)
    add_bullet(doc, "Results apply to the tested public benchmark versions, prompts, sampling settings, endpoints, and evaluator toolchain.")
    add_bullet(doc, "No prompt engineering or automated repair loop is included in the baseline condition.")
    add_bullet(doc, "Pass@5 measures whether at least one of five samples solves a task; it is not the same as single-sample reliability.")
    add_bullet(doc, "ProtocolLLM public evaluation is lint-only in the current harness.")
    add_bullet(doc, "RTL-OPT generic cell counts are an early proxy, not final PPA evidence from a technology-mapped flow.")
    add_bullet(doc, "One DeepSeek VerilogEval pass@5 attempt was incomplete and excluded; one ProtocolLLM attempt failed due to endpoint timeout and was rerun successfully.")
    add_heading(doc, "Appendix C. Reproducibility", 1)
    add_body(doc, "The benchmark repository records model presets, deterministic benchmark configuration, timestamped output paths, failure categories, and comparison reports. Raw output folders remain on Lanta; Git tracks the configs, summary reports, and output manifests.")

    core = doc.core_properties
    core.title = "Evaluating Large Language Models for RTL Generation"
    core.subject = "Public RTL benchmark comparison"
    core.author = "SiliconCraft RTL Benchmarking Project"
    core.keywords = "RTL, Verilog, LLM, benchmark, VerilogEval, RTLLM, RTL-OPT"
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
